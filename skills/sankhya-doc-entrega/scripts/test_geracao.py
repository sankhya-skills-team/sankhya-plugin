"""Auto-teste dos geradores. Executar: python scripts/test_geracao.py

Gera HTML e DOCX de exemplo em um diretorio temporario e verifica
versionamento, backup, historico e presenca do logo. Sem framework.
"""

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile

AQUI = os.path.dirname(os.path.abspath(__file__))

EXEMPLO = {
    "parceiro": "Parceiro Teste",
    "id_demanda": "DEM-0001",
    "nome_customizacao": "Pesagem de Entrada",
    "caminho_sistema": "Menu › Beneficiamento › BEN — Pesagem de Entrada",
    "responsavel_tecnico": "Dev A → Dev B (a partir da v1.1)",
    "objetivo": "Automatiza o registro de pesagem & libera a nota <fiscal>.",
    "limitacoes_gerais": ["Não reprocessa pesagens já encerradas."],
    "incluir_homologacao": True,
    "incluir_assinaturas": True,
    "personas_sankhya": ["Ana Paula Souza — Gerente de Projetos"],
    "personas_cliente": ["Roberto Mendes — Diretor Comercial"],
    "funcionalidades": [{
        "titulo": "Calcular Pesagem",
        "tipo": "acao",
        "icone": "⚖️",
        "passos": ["O usuário seleciona o ticket.", "O sistema calcula o peso líquido."],
        "obs": "Requer perfil Balança.",
        "limitacoes": "Irreversível após o encerramento.",
        "tipo_acesso": "tela",
        "testes": [
            {"nome": "Executar com ticket em aberto",
             "esperado": "Operação concluída — peso líquido gravado"},
            {"nome": "Executar com ticket encerrado",
             "esperado": 'Sistema bloqueia com mensagem: "Ticket já encerrado."'},
        ],
    }],
    "checklist_deploy": {
        "pre_requisitos": [
            {"tipo": "tela_adicional", "nome": "AD_PESAGEM",
             "arquivo": "Metadados_AD_PESAGEM.zip", "observacao": ""},
            {"tipo": "parametro", "nome": "PESAGEM_TOL", "descricao": "Tolerância",
             "tipo_valor": "número", "valor_padrao": "0.5"},
        ],
        "pos_deploy": [
            {"tipo": "acao", "nome_exibicao": "Calcular Pesagem", "entidade": "AD_PESAGEM",
             "tipo_sankhya": "AcaoRotinaJava", "classe": "br.com.sankhya.X", "perfis": "Balança"},
            {"tipo": "jar", "arquivo": "pesagem-1.0.0.jar", "caminho_servidor": ""},
        ],
    },
}


def gerar(script, dados, destino):
    dados = dict(dados, arquivo_saida=destino)
    os.makedirs(os.path.dirname(destino), exist_ok=True)
    entrada = os.path.join(os.path.dirname(destino), "_dados.json")
    with open(entrada, "w", encoding="utf-8") as f:
        json.dump(dados, f, ensure_ascii=False)
    saida = subprocess.run([sys.executable, os.path.join(AQUI, script), entrada],
                           capture_output=True, text=True, encoding="utf-8")
    assert saida.returncode == 0, "%s falhou:\n%s" % (script, saida.stderr)
    return json.loads(saida.stdout.strip().splitlines()[-1])


def main():
    tmp = tempfile.mkdtemp(prefix="doc-entrega-")
    try:
        # ── HTML ───────────────────────────────────────────────────
        alvo = os.path.join(tmp, "Documentacao", "Entrega - Pesagem de Entrada.html")
        r1 = gerar("gerar_html.py", EXEMPLO, alvo)
        assert r1["versao"] == "1.0" and r1["backup"] is None, r1
        html = open(alvo, encoding="utf-8").read()

        assert "&lt;fiscal&gt;" in html and "&amp;" in html, "escape HTML quebrado"
        assert "<svg" in html and "currentColor" in html, "logo SVG ausente"
        assert "#4ADE80" in html and "#243143" in html, "paleta do design system ausente"
        assert "#6AA84F" not in html and "#1a4d2e" not in html, "cor antiga remanescente"
        assert html.count("class=\"crumb") == 3, "trilha de navegação incorreta"
        assert "congelarEstado" in html, "correção do export ausente"
        assert html.count("pesagem-1.0.0.jar") == 1, "nome do JAR repetido no detalhe"
        assert "Histórico de versões" in html
        assert re.search(r'name="doc-versao" content="1\.0"', html)
        # nenhum placeholder de formatação sobrou no CSS
        assert "%(" not in html.split("<script>")[0], "placeholder %( no HTML/CSS"

        r2 = gerar("gerar_html.py", dict(EXEMPLO, changelog=["Ajuste de tolerância."]), alvo)
        assert r2["versao"] == "1.1", r2
        assert r2["backup"] and os.path.exists(r2["backup"]), "backup não criado"
        html2 = open(alvo, encoding="utf-8").read()
        assert "Ajuste de tolerância." in html2 and "v1.0" in html2, "histórico não acumulou"

        # ── DOCX ───────────────────────────────────────────────────
        alvo_dx = os.path.join(tmp, "Documentacao", "Entrega - Pesagem de Entrada.docx")
        r3 = gerar("gerar_docx.py", EXEMPLO, alvo_dx)
        assert r3["versao"] == "1.0", r3          # extensao diferente = versao propria
        r4 = gerar("gerar_docx.py", dict(EXEMPLO, changelog=["Revisão do escopo."]), alvo_dx)
        assert r4["versao"] == "1.1" and os.path.exists(r4["backup"]), r4

        from docx import Document
        doc = Document(alvo_dx)
        assert doc.core_properties.version == "1.1"
        texto = "\n".join(p.text for p in doc.paragraphs)

        # ABNT NBR 14724: margens 3/2/3/2 cm e entrelinha 1,5 no corpo.
        # O Word grava em twips, então a volta em cm não bate no EMU exato.
        s = doc.sections[0]
        margens = [m.cm for m in (s.top_margin, s.bottom_margin,
                                  s.left_margin, s.right_margin)]
        assert all(abs(m - alvo) < 0.01 for m, alvo in zip(margens, (3, 2, 3, 2))), \
            "margens fora da ABNT: %s" % margens
        assert doc.styles["Normal"].paragraph_format.line_spacing == 1.5, "entrelinha fora da ABNT"
        assert doc.styles["Normal"].font.size.pt == 12, "corpo fora da ABNT"
        assert all(p.paragraph_format.line_spacing == 1.0
                   for t in doc.tables for r in t.rows for c in r.cells
                   for p in c.paragraphs), "tabela deveria usar espaço simples"
        # A numeração dos passos reinicia a cada funcionalidade.
        assert texto.count("1. O usuário seleciona o ticket.") == 1
        assert "Entrega de Desenvolvimento" in texto and "Checklist de Deploy" in texto
        assert "Homologação e Testes" in texto
        assert len(doc.inline_shapes) == 1, "logo ausente no DOCX"
        assert doc.element.body.xml.count('w:fill="243143"') >= 4, "cabeçalhos sem shading"

        # Sem layout fixo o Word ignora as larguras: o autofit entrega a faixa
        # à coluna de texto mais longo e espreme as outras (a coluna "OK" do
        # checklist quebrava "[ ]" em duas linhas).
        secao = doc.sections[0]
        util = secao.page_width - secao.left_margin - secao.right_margin
        assert doc.tables, "documento sem tabelas"
        for n, tabela in enumerate(doc.tables):
            assert 'w:type="fixed"' in tabela._tbl.tblPr.xml, "tabela %d sem layout fixo" % n
            for linha in tabela.rows:
                # Numa linha com merge, cells[] repete a mesma célula em cada
                # índice que ela cobre; contar duas vezes dobraria a soma.
                unicas, vistos = [], set()
                for celula in linha.cells:
                    if id(celula._tc) not in vistos:
                        vistos.add(id(celula._tc))
                        unicas.append(celula)
                larguras = [c.width for c in unicas]
                assert all(larguras), "tabela %d tem célula sem largura" % n
                assert sum(larguras) <= util, "tabela %d estoura a faixa útil" % n

        print("OK — HTML e DOCX gerados, versionados e validados.")
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    main()
