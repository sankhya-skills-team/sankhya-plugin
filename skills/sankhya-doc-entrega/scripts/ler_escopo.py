"""Extrai o texto de um documento de escopo (.md, .txt, .docx, .pdf).

Uso: python ler_escopo.py <caminho> [limite_de_caracteres]
Imprime o texto extraido na saida padrao.
"""

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from _comum import garantir


def ler_docx(caminho):
    garantir("docx", "python-docx")
    from docx import Document
    doc = Document(caminho)
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabela in doc.tables:
        for linha in tabela.rows:
            celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
            if celulas:
                partes.append(" | ".join(celulas))
    return "\n".join(partes)


def ler_pdf(caminho):
    garantir("pdfplumber")
    import pdfplumber
    with pdfplumber.open(caminho) as pdf:
        return "\n".join(p.extract_text() or "" for p in pdf.pages)


def ler_escopo(caminho):
    ext = os.path.splitext(caminho)[1].lower()
    if ext in (".md", ".txt"):
        with open(caminho, encoding="utf-8", errors="replace") as f:
            return f.read()
    if ext in (".docx", ".doc"):
        return ler_docx(caminho)
    if ext == ".pdf":
        return ler_pdf(caminho)
    raise SystemExit("extensao nao suportada: %s" % ext)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        raise SystemExit("uso: python ler_escopo.py <caminho> [limite]")
    limite = int(sys.argv[2]) if len(sys.argv) > 2 else 0
    texto = ler_escopo(sys.argv[1]).strip()
    if limite and len(texto) > limite:
        texto = texto[:limite] + "\n[... truncado ...]"
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    print(texto)
