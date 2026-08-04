"""Gera o documento de entrega em .docx (Word editavel).

Uso: python gerar_docx.py <dados.json>
Contrato do JSON: ver secao "Contrato dados.json" no SKILL.md.
Imprime na saida padrao um JSON com {arquivo, versao, backup}.

Regra visual: usa apenas as cores do design system Sankhya e nao aplica
cor de fundo na pagina — o preenchimento fica restrito a cabecalhos de tabela.
"""

import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

import _brand as B
from _comum import carregar_dados, garantir, parse_personas, resolver_versao

garantir("docx", "python-docx")

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

D = carregar_dados(sys.argv)

ARQUIVO_SAIDA = D["arquivo_saida"]
VERSAO, BACKUP, HISTORICO = resolver_versao(ARQUIVO_SAIDA, D.get("changelog"))
DATA_GERACAO = HISTORICO[0]["data"]

FUNCIONALIDADES = D.get("funcionalidades", [])
CHECKLIST       = D.get("checklist_deploy", {}) or {}
PRE  = CHECKLIST.get("pre_requisitos") or []
POS  = CHECKLIST.get("pos_deploy") or []

C_TERTIARY  = RGBColor(*B.rgb(B.TERTIARY))
C_ON_SURF   = RGBColor(*B.rgb(B.ON_SURFACE))
C_WHITE     = RGBColor(255, 255, 255)
HEX_TERTIARY = B.hexr(B.TERTIARY)


# ── ABNT NBR 14724 ─────────────────────────────────────────────────
# 5.1 margens: 3 cm em cima e a esquerda, 2 cm embaixo e a direita.
# 5.2 espacamento: 1,5 entre linhas no texto; simples em tabelas,
#     legendas e notas, nestas com corpo menor.
ABNT_MARGEM_SUP  = Cm(3)
ABNT_MARGEM_INF  = Cm(2)
ABNT_MARGEM_ESQ  = Cm(3)
ABNT_MARGEM_DIR  = Cm(2)
ABNT_CORPO       = Pt(12)
ABNT_ENTRELINHA  = 1.5
ABNT_LINHA       = Pt(18)   # uma entrelinha de 1,5 sobre corpo de 12 pt
ABNT_TABELA      = 1.0      # espaco simples dentro de tabela

ESTILOS_TEXTO = ("Normal", "List Bullet", "List Number")


def aplicar_abnt_paginas(doc):
    for secao in doc.sections:
        secao.top_margin = ABNT_MARGEM_SUP
        secao.bottom_margin = ABNT_MARGEM_INF
        secao.left_margin = ABNT_MARGEM_ESQ
        secao.right_margin = ABNT_MARGEM_DIR


def aplicar_abnt_estilos(doc):
    for nome in ESTILOS_TEXTO:
        try:
            estilo = doc.styles[nome]
        except KeyError:
            continue
        estilo.font.size = ABNT_CORPO
        pf = estilo.paragraph_format
        pf.line_spacing = ABNT_ENTRELINHA
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)


def aplicar_abnt_tabelas(doc):
    """Espaco simples dentro das tabelas, como manda a norma.

    Roda no fim: as celulas herdam o estilo Normal, que ja esta em 1,5.
    Mexe so na entrelinha -- o espacamento vertical entre paragrafos ja vem
    zerado do estilo, e zerar de novo apagaria o vao entre as assinaturas.
    """
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for par in celula.paragraphs:
                    par.paragraph_format.line_spacing = ABNT_TABELA


# ── Helpers de formatacao ──────────────────────────────────────────

def sombrear(celula, hex_fill):
    tcPr = celula._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def run(par, texto, bold=None, size=None, color=None, italic=None):
    r = par.add_run(texto)
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return r


def titulo(doc, texto, nivel=2):
    p = doc.add_paragraph()
    run(p, texto, bold=True, size=14 if nivel == 2 else 12, color=C_TERTIARY)
    # ABNT 5.2: titulo separado do texto anterior e posterior por uma
    # entrelinha de 1,5 -- 18 pt sobre corpo de 12 pt.
    p.paragraph_format.space_before = ABNT_LINHA
    p.paragraph_format.space_after = ABNT_LINHA
    p.paragraph_format.line_spacing = 1.0
    if nivel == 2:
        _borda_inferior(p, B.hexr(B.PRIMARY))
    return p


def _borda_inferior(par, hex_cor):
    pPr = par._p.get_or_add_pPr()
    bordas = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), hex_cor)
    bordas.append(bottom)
    pPr.append(bordas)


def cabecalho_tabela(tabela, colunas):
    for i, texto in enumerate(colunas):
        celula = tabela.rows[0].cells[i]
        sombrear(celula, HEX_TERTIARY)
        run(celula.paragraphs[0], texto, bold=True, size=10, color=C_WHITE)


def linhas(doc, valor, estilo="List Bullet"):
    """Aceita string (com \\n) ou lista de strings."""
    itens = valor if isinstance(valor, list) else str(valor or "").split("\n")
    for item in itens:
        texto = str(item).strip()
        if texto:
            doc.add_paragraph(texto, style=estilo)


def linhas_numeradas(doc, valor):
    """Passos numerados no proprio texto, com recuo deslocado.

    Nao usa o estilo List Number: todos os paragrafos dele compartilham a
    mesma instancia de numeracao, entao a contagem seguia acumulando de uma
    funcionalidade para a proxima (2.2 comecava no 7).
    """
    itens = valor if isinstance(valor, list) else str(valor or "").split("\n")
    itens = [str(i).strip() for i in itens if str(i).strip()]
    for n, texto in enumerate(itens, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        run(p, "%d. " % n, bold=True, color=C_TERTIARY)
        run(p, texto)


def layout_fixo(tabela):
    """Faz o Word respeitar as larguras declaradas.

    Sem w:tblLayout fixed o Word roda o autofit e ignora tanto
    table.autofit=False quanto os valores de columns[i].width.
    """
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tabela._tbl.tblPr.append(layout)


def celula_assinatura(celula, papel, ultima=False):
    # Borda inferior em vez de underscores: acompanha a largura da celula
    # em qualquer fonte, sem estourar para a linha seguinte.
    p1 = celula.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    _borda_inferior(p1, "808080")
    p2 = celula.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    run(p2, "Nome:", bold=True)
    p3 = celula.add_paragraph()
    p3.paragraph_format.space_after = Pt(0) if ultima else Pt(36)
    run(p3, papel, color=C_ON_SURF)


# ── Documento ──────────────────────────────────────────────────────

doc = Document()
for estilo in doc.styles:
    if hasattr(estilo, "font"):
        estilo.font.name = B.FONT_DOCX
aplicar_abnt_paginas(doc)
aplicar_abnt_estilos(doc)

# Cabecalho com logo
if os.path.exists(B.LOGO_PNG):
    p_logo = doc.add_paragraph()
    p_logo.paragraph_format.space_after = ABNT_LINHA
    p_logo.add_run().add_picture(B.LOGO_PNG, width=Cm(4.0))

p_tit = doc.add_paragraph()
run(p_tit, "Entrega de Desenvolvimento", bold=True, size=18, color=C_TERTIARY)
p_sub = doc.add_paragraph()
run(p_sub, D.get("nome_customizacao", ""), size=13, color=C_ON_SURF)
run(p_sub, "   v%s" % VERSAO, bold=True, size=13, color=RGBColor(*B.rgb(B.PRIMARY)))
p_sub.paragraph_format.space_after = ABNT_LINHA

# ── Identificacao ──────────────────────────────────────────────────
dados_id = [
    ("Parceiro", D.get("parceiro", "")),
    ("ID Demanda", D.get("id_demanda") or "—"),
    ("Responsável Técnico", D.get("responsavel_tecnico", "")),
    ("Versão da Customização", VERSAO),
    ("Data de Geração", DATA_GERACAO),
    ("Caminho no Sistema", D.get("caminho_sistema", "")),
]
tab_id = doc.add_table(rows=1 + len(dados_id), cols=2)
tab_id.style = "Table Grid"
celula_cab = tab_id.rows[0].cells[0].merge(tab_id.rows[0].cells[1])
sombrear(celula_cab, HEX_TERTIARY)
run(celula_cab.paragraphs[0], "Identificação", bold=True, size=11, color=C_WHITE)
for i, (rotulo, valor) in enumerate(dados_id, 1):
    linha = tab_id.rows[i]
    run(linha.cells[0].paragraphs[0], rotulo, bold=True, color=C_TERTIARY)
    run(linha.cells[1].paragraphs[0], str(valor), color=C_ON_SURF)

# ── Historico de versoes ───────────────────────────────────────────
titulo(doc, "Histórico de versões", nivel=3)
tab_h = doc.add_table(rows=1 + len(HISTORICO), cols=3)
tab_h.style = "Table Grid"
cabecalho_tabela(tab_h, ["Versão", "Data", "Alterações"])
for i, entrada in enumerate(HISTORICO, 1):
    linha = tab_h.rows[i]
    run(linha.cells[0].paragraphs[0], "v" + entrada["versao"], bold=True, color=C_TERTIARY)
    run(linha.cells[1].paragraphs[0], entrada["data"], color=C_ON_SURF)
    alteracoes = entrada.get("alteracoes") or ["—"]
    celula = linha.cells[2]
    run(celula.paragraphs[0], alteracoes[0], color=C_ON_SURF)
    for extra in alteracoes[1:]:
        run(celula.add_paragraph(), extra, color=C_ON_SURF)

# ── 1. Objetivos ───────────────────────────────────────────────────
titulo(doc, "1. Objetivos")
doc.add_paragraph(D.get("objetivo", ""))

# ── 2. Manual de uso ───────────────────────────────────────────────
titulo(doc, "2. Manual de Uso das Customizações")
p_c = doc.add_paragraph()
run(p_c, "Caminho no sistema: ", bold=True, color=C_TERTIARY)
run(p_c, D.get("caminho_sistema", ""), color=C_ON_SURF)

for i, func in enumerate(FUNCIONALIDADES, 1):
    meta = B.TIPO_META.get(func.get("tipo"), B.TIPO_META["acao"])
    titulo(doc, "2.%d %s" % (i, func.get("titulo", "")), nivel=3)
    p_t = doc.add_paragraph()
    run(p_t, meta["badge"], bold=True, size=9, color=RGBColor(*B.rgb(meta["cor"])))
    p_p = doc.add_paragraph()
    p_p.paragraph_format.space_before = ABNT_LINHA
    run(p_p, "Passo a passo:", bold=True, color=C_TERTIARY)
    linhas_numeradas(doc, func.get("passos", []))
    if func.get("obs"):
        p_o = doc.add_paragraph()
        p_o.paragraph_format.space_before = ABNT_LINHA
        run(p_o, "Observações: ", bold=True, color=C_TERTIARY)
        run(p_o, str(func["obs"]), color=C_ON_SURF)
    if func.get("limitacoes"):
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = ABNT_LINHA
        run(p_l, "Limitações: ", bold=True, color=C_TERTIARY)
        run(p_l, str(func["limitacoes"]), color=C_ON_SURF)

# ── 3. Limitacoes conhecidas ───────────────────────────────────────
titulo(doc, "3. Limitações Conhecidas")
lims = D.get("limitacoes_gerais") or []
if lims:
    linhas(doc, lims)
else:
    run(doc.add_paragraph(), "Nenhuma limitação global identificada.",
        italic=True, color=C_ON_SURF)

secao_n = 4

# ── 4. Checklist de deploy ─────────────────────────────────────────
LABEL_TIPO = {"tela_adicional": "Tela Adicional", "parametro": "Parâmetro TSIPAR",
              "script_sql": "Script DDL", "acao": "Botão de Ação", "evento": "Evento",
              "job": "Job", "regra": "Regra", "jar": "Deploy JAR",
              "dashboard": "Dashboard", "relatorio": "Relatório"}
CAMPOS_DETALHE = [("arquivo", "Arquivo"), ("entidade", "Entidade"),
                  ("tipo_sankhya", "Tipo"), ("classe", "Classe"),
                  ("perfis", "Perfis"), ("descricao", None),
                  ("tipo_valor", "Tipo"), ("valor_padrao", "Padrão"),
                  ("caminho_servidor", "Destino"), ("observacao", None)]


def detalhe_item(item):
    nome = item.get("nome_exibicao") or item.get("nome") or item.get("arquivo", "")
    partes = []
    for chave, rotulo in CAMPOS_DETALHE:
        valor = item.get(chave)
        # Sem nome proprio, o arquivo virou o titulo: nao repetir no detalhe.
        if valor and valor != nome:
            partes.append("%s: %s" % (rotulo, valor) if rotulo else str(valor))
    return " | ".join(partes)


def tabela_checklist(doc, itens):
    tab = doc.add_table(rows=1 + len(itens), cols=4)
    tab.style = "Table Grid"
    cabecalho_tabela(tab, ["OK", "Tipo", "Item", "Detalhes"])
    tab.columns[0].width = Inches(0.4)
    for i, item in enumerate(itens, 1):
        linha = tab.rows[i]
        run(linha.cells[0].paragraphs[0], "[  ]", color=C_ON_SURF)
        run(linha.cells[1].paragraphs[0],
            LABEL_TIPO.get(item.get("tipo", ""), item.get("tipo", "")),
            size=9, color=C_ON_SURF)
        nome = item.get("nome_exibicao") or item.get("nome") or item.get("arquivo", "")
        run(linha.cells[2].paragraphs[0], str(nome), bold=True, color=C_TERTIARY)
        run(linha.cells[3].paragraphs[0], detalhe_item(item), size=9, color=C_ON_SURF)


if PRE or POS:
    titulo(doc, "%d. Checklist de Deploy" % secao_n)
    if PRE:
        titulo(doc, "%d.1 Pré-requisitos — fazer ANTES do deploy" % secao_n, nivel=3)
        tabela_checklist(doc, PRE)
    if POS:
        titulo(doc, "%d.%d Pós-deploy — configurar APÓS o JAR no servidor"
               % (secao_n, 2 if PRE else 1), nivel=3)
        tabela_checklist(doc, POS)
    secao_n += 1

# ── 5. Homologacao ─────────────────────────────────────────────────
titulo(doc, "%d. Homologação e Testes" % secao_n)
testes_por_func = [(f, f.get("testes") or []) for f in FUNCIONALIDADES]
tem_testes = any(t for _f, t in testes_por_func)

if tem_testes:
    p_i = doc.add_paragraph()
    run(p_i, "Marque o resultado de cada cenário e anexe as evidências ao final "
             "do documento.", italic=True, size=10, color=C_ON_SURF)
    for func, testes in testes_por_func:
        if not testes:
            continue
        titulo(doc, func.get("titulo", ""), nivel=3)
        tab = doc.add_table(rows=1 + len(testes), cols=3)
        tab.style = "Table Grid"
        cabecalho_tabela(tab, ["Cenário", "Resultado esperado", "Resultado"])
        for i, teste in enumerate(testes, 1):
            linha = tab.rows[i]
            run(linha.cells[0].paragraphs[0], teste.get("nome", ""), color=C_TERTIARY)
            run(linha.cells[1].paragraphs[0], teste.get("esperado", ""), color=C_ON_SURF)
            run(linha.cells[2].paragraphs[0], "[ ] Aprovado\n[ ] Reprovado",
                size=9, color=C_ON_SURF)

for i, rotulo in enumerate(
        ("Responsável pelos testes no cliente: ______________________________",
         "Data da homologação: ____/____/________",
         "Resultado geral: [  ] Aprovado    [  ] Reprovado (pendências abaixo)")):
    p = doc.add_paragraph(rotulo)
    if i == 0:
        p.paragraph_format.space_before = ABNT_LINHA
secao_n += 1

# ── Observacoes ────────────────────────────────────────────────────
titulo(doc, "%d. Observações" % secao_n)
doc.add_paragraph(
    "Por se tratar de uma personalização o Service Desk não atua nos casos de dúvidas "
    "ou incidentes; qualquer alteração ou suporte deverá ser negociada como horas "
    "adicionais com a Unidade para o atendimento.")

# ── Assinaturas ────────────────────────────────────────────────────
if D.get("incluir_assinaturas", True):
    p_cd = doc.add_paragraph("_______________, ____ de ______________ de ______")
    p_cd.paragraph_format.space_before = Pt(36)
    p_cd.paragraph_format.space_after = Pt(24)

    sankhya = parse_personas(D.get("personas_sankhya")) or [
        {"nome": "", "funcao": "Consultor"},
        {"nome": "", "funcao": "Gerente de Projetos – Sankhya"}]
    cliente = parse_personas(D.get("personas_cliente")) or [
        {"nome": "", "funcao": "Líder do Projeto"},
        {"nome": "", "funcao": "Solicitante"}]

    pares = max(len(sankhya), len(cliente))
    tab_s = doc.add_table(rows=pares, cols=3)
    tab_s.autofit = False
    layout_fixo(tab_s)

    # Larguras derivadas da faixa util da pagina para sobreviver a mudanca
    # de margem. A coluna do meio e so o vao entre as duas assinaturas.
    secao = doc.sections[0]
    util = secao.page_width - secao.left_margin - secao.right_margin
    vao = int(util * 0.12)
    coluna = (util - vao) // 2
    larguras = (coluna, vao, coluna)

    for c, largura in enumerate(larguras):
        tab_s.columns[c].width = largura
    for linha in tab_s.rows:
        # A largura so vale se estiver em cada celula, inclusive na do vao.
        for c, largura in enumerate(larguras):
            linha.cells[c].width = largura

    for i in range(pares):
        ultima = (i == pares - 1)
        if i < len(sankhya):
            p = sankhya[i]
            celula_assinatura(tab_s.cell(i, 0),
                              p["funcao"] or p["nome"] or "Sankhya", ultima)
        if i < len(cliente):
            p = cliente[i]
            celula_assinatura(tab_s.cell(i, 2),
                              p["funcao"] or p["nome"] or "Cliente", ultima)

aplicar_abnt_tabelas(doc)

doc.core_properties.version = VERSAO
doc.core_properties.title = "Entrega — %s" % D.get("nome_customizacao", "")
doc.save(ARQUIVO_SAIDA)

print(json.dumps({"arquivo": ARQUIVO_SAIDA, "versao": VERSAO, "backup": BACKUP}))
