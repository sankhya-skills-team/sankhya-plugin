#!/usr/bin/env python3
import argparse
from pathlib import Path

from docx import Document


def esc(text):
    return (text or "").replace("|", "\\|").strip()


def para_to_md(paragraph):
    text = paragraph.text.strip()
    if not text:
        return ""
    style = paragraph.style.name if paragraph.style is not None else ""
    if style.startswith("Heading "):
        try:
            level = int(style.split()[-1])
        except ValueError:
            level = 1
        return f"{'#' * max(1, min(level, 6))} {text}"
    if style.startswith("List Bullet"):
        return f"- {text}"
    if style.startswith("List Number"):
        return f"1. {text}"
    return text


def table_to_md(table):
    rows = []
    for row in table.rows:
        rows.append([esc(cell.text.replace("\n", " ")) for cell in row.cells])
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    out = []
    out.append("| " + " | ".join(rows[0]) + " |")
    out.append("| " + " | ".join(["---"] * width) + " |")
    for row in rows[1:]:
        out.append("| " + " | ".join(row) + " |")
    return "\n".join(out)


def iter_blocks(doc):
    body = doc.element.body
    para_map = {p._p: p for p in doc.paragraphs}
    table_map = {t._tbl: t for t in doc.tables}
    for child in body.iterchildren():
        if child in para_map:
            md = para_to_md(para_map[child])
            if md:
                yield md
        elif child in table_map:
            md = table_to_md(table_map[child])
            if md:
                yield md


def main():
    ap = argparse.ArgumentParser(description="Fallback DOCX to Markdown extractor using python-docx.")
    ap.add_argument("docx", help="Input .docx file")
    ap.add_argument("-o", "--output", help="Output .md file. Defaults to stdout.")
    args = ap.parse_args()

    doc = Document(args.docx)
    text = "\n\n".join(iter_blocks(doc)) + "\n"
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
    else:
        print(text, end="")


if __name__ == "__main__":
    main()
