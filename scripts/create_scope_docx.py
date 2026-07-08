#!/usr/bin/env python3
import argparse
import json
from copy import deepcopy
from pathlib import Path

from docx import Document


SKILL_DIR = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_DIR / "assets" / "escopo-template.docx"


def clear_body(doc):
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    return body


def add_p(doc, text="", style="Normal", bold=False):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        run.bold = bold
    return p


def add_heading(doc, text, level=1):
    return add_p(doc, text, f"Heading {level}")


def add_bullet(doc, text):
    return add_p(doc, text, "List Bullet")


def add_numbered(doc, items):
    # Textual numbering avoids cross-section numbering continuation in inherited templates.
    for idx, text in enumerate(items, 1):
        add_p(doc, f"{idx}. {text}", "Normal")


def add_table(doc, headers, rows, template_table=None):
    table = doc.add_table(rows=1, cols=len(headers))
    if template_table is not None:
        table.style = template_table.style
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    add_p(doc, "")
    return table


def update_identification(table, values):
    if not values:
        return
    existing = {}
    for row in table.rows:
        if len(row.cells) >= 2:
            key = row.cells[0].text.strip().rstrip(":")
            existing[key] = row
    for key, value in values.items():
        row = existing.get(key)
        if row is None:
            row = table.add_row()
            row.cells[0].text = f"{key}:"
        row.cells[1].text = str(value)


def render_section(doc, section, table_style_source=None):
    heading = section.get("heading")
    if heading:
        level = int(section.get("level", 1))
        add_heading(doc, heading, level)
    for text in section.get("paragraphs", []):
        add_p(doc, text)
    if section.get("numbered"):
        add_numbered(doc, section["numbered"])
    for text in section.get("bullets", []):
        add_bullet(doc, text)
    for tbl in section.get("tables", []):
        add_table(doc, tbl["headers"], tbl["rows"], table_style_source)


def main():
    ap = argparse.ArgumentParser(description="Generate scope DOCX preserving a template's Word styles.")
    ap.add_argument("--content", required=True, help="JSON content file.")
    ap.add_argument("--output", required=True, help="Output .docx path.")
    ap.add_argument("--template", default=str(DEFAULT_TEMPLATE), help="Template .docx path.")
    args = ap.parse_args()

    content = json.loads(Path(args.content).read_text(encoding="utf-8"))
    template = Path(args.template)
    output = Path(args.output)

    doc = Document(template)
    base = Document(template)
    ident = deepcopy(base.tables[0]._tbl) if base.tables else None
    signature = deepcopy(base.tables[-1]._tbl) if len(base.tables) >= 2 else None
    table_style_source = base.tables[1] if len(base.tables) >= 2 else None

    body = clear_body(doc)
    if ident is not None:
        body.insert(0, ident)
        update_identification(doc.tables[0], content.get("identification", {}))

    subtitle = content.get("subtitle") or f"Proposta Técnica: {content.get('title', 'Escopo')}"
    add_p(doc, subtitle, "Normal", True)

    for section in content.get("sections", []):
        render_section(doc, section, table_style_source)

    for table_def in content.get("tables", []):
        heading = table_def.get("after_heading")
        if heading:
            add_heading(doc, heading, int(table_def.get("level", 2)))
        add_table(doc, table_def["headers"], table_def["rows"], table_style_source)

    if signature is not None:
        body = doc._body._element
        body.insert(len(body) - 1, signature)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    main()
